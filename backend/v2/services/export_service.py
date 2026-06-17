"""Export service — convert OCR results to various formats."""

from __future__ import annotations

import csv
import io
import json
from pathlib import Path
from typing import Optional

from config import EXPORT_DIR


class ExportService:
    """Convert OCR results to JSON, TXT, CSV, DOCX, or searchable PDF."""

    @staticmethod
    def export_json(result_dict: dict) -> str:
        """Export result as formatted JSON string."""
        return json.dumps(result_dict, indent=2, ensure_ascii=False)

    @staticmethod
    def export_txt(result_dict: dict) -> str:
        """Extract just the full text from result dict."""
        return result_dict.get("full_text", "")

    @staticmethod
    def export_csv(result_dict: dict) -> str:
        """Export lines as CSV string."""
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(["index", "text", "confidence", "bounding_box"])
        for line in result_dict.get("lines", []):
            writer.writerow([
                line["index"],
                line["text"],
                line.get("confidence", ""),
                ",".join(str(v) for v in line.get("bounding_box", [])),
            ])
        return output.getvalue()

    @staticmethod
    def export_docx(result_dict: dict) -> bytes:
        """Export result as a .docx document."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required for DOCX export. Install with: pip install python-docx")

        doc = Document()
        doc.add_heading(f"OCR Result: {result_dict.get('filename', 'document')}", level=1)

        stats = result_dict.get("confidence_stats", {})
        if stats:
            doc.add_paragraph(f"Avg Confidence: {stats.get('mean', 'N/A')}")

        doc.add_paragraph(f"Detected Lines: {result_dict.get('detected_lines', 0)}")
        doc.add_paragraph(f"Processing Time: {result_dict.get('processing_time_ms', 0):.1f}ms")
        doc.add_paragraph("")

        full_text = result_dict.get("full_text", "")
        if full_text:
            doc.add_heading("Full Text", level=2)
            doc.add_paragraph(full_text)

        lines = result_dict.get("lines", [])
        if lines:
            doc.add_heading("Line Details", level=2)
            table = doc.add_table(rows=1, cols=4)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Index"
            hdr_cells[1].text = "Text"
            hdr_cells[2].text = "Confidence"
            hdr_cells[3].text = "Bounding Box"

            for line in lines:
                row_cells = table.add_row().cells
                row_cells[0].text = str(line["index"])
                row_cells[1].text = line["text"]
                row_cells[2].text = str(line.get("confidence", ""))
                bbox = line.get("bounding_box", [])
                row_cells[3].text = f"[{bbox[0]:.0f}, {bbox[1]:.0f}, {bbox[2]:.0f}, {bbox[3]:.0f}]" if bbox else ""

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    @staticmethod
    def export_searchable_pdf(result_dict: dict, output_dir: Optional[Path] = None) -> bytes:
        """Export OCR result as a searchable PDF with embedded text layer.

        Requires fitz (PyMuPDF) and the annotated image.
        This creates a PDF page with the annotated image on it and invisible text overlay.
        """
        import base64
        import fitz
        from PIL import Image, ImageDraw, ImageFont

        buf = io.BytesIO()
        doc = fitz.open()

        # Get or decode annotated image
        b64_img = result_dict.get("annotated_image_b64")
        if b64_img:
            img_bytes = base64.b64decode(b64_img)
            img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
        else:
            # Create blank white image if no annotated image available
            full_text = result_dict.get("full_text", "No text detected")
            lines = full_text.split("\n") if full_text else []
            line_height = 30
            h = max(len(lines) * line_height + 100, 400)
            img = Image.new("RGB", (612, h), "white")
            draw = ImageDraw.Draw(img)
            try:
                font = ImageFont.truetype("arial.ttf", 14)
            except OSError:
                font = ImageFont.load_default()
            y = 50
            for line in lines:
                draw.text((50, y), line[:80], fill="black", font=font)
                y += line_height

        # Convert PIL to fitz pixmap
        img_array = list(img.getdata())
        width, height = img.size
        pixmap = fitz.Pixmap(fitz.csRGB, 0, 0, 0, width, height)
        pixmap.set_image(data=img_array)

        page = doc.new_page(width=fitz.paper_size("a4")[0], height=height)
        # Insert image
        rect = fitz.Rect(0, 0, width, height)
        page.insert_image(rect, pixmap=pixmap)

        # Add invisible text layer for searchability
        if result_dict.get("lines"):
            font_size = 14
            try:
                font = fitz.Font(font="urdu")
            except Exception:
                font = None

            for line_result in result_dict["lines"]:
                bbox = line_result.get("bounding_box", [])
                if len(bbox) >= 4 and line_result.get("text"):
                    text_rect = fitz.Rect(bbox[0], bbox[1], bbox[2], bbox[1] + font_size)
                    page.insert_text(
                        text_rect.tl,
                        line_result["text"],
                        fontsize=font_size,
                        fontname=font.name if font else None,
                        color=(0, 0, 0),
                        opacity=0,  # invisible but searchable
                    )

        pdf_bytes = bytearray()
        doc.save(pdf_bytes, garbage=4, deflate=True)
        doc.close()
        return bytes(pdf_bytes)

    @staticmethod
    def save_export(export_data: bytes, filename: str, output_dir: Optional[Path] = None) -> Path:
        """Save export data to disk. Returns the file path."""
        dest = output_dir or EXPORT_DIR
        dest.mkdir(parents=True, exist_ok=True)
        path = dest / filename
        with open(path, "wb") as f:
            f.write(export_data)
        return path
