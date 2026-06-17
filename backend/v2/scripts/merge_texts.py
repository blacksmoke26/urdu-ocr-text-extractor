#!/usr/bin/env python3
"""
merge_texts_v2.py — Merge multiple OCR text outputs into a single formatted document.

Usage:
    python scripts/merge_texts_v2.py -i ./results/page_0001/text.txt ./results/page_0002/text.txt ... -o merged.txt
    python scripts/merge_texts_v2.py --dir ./ocr_output -o doc.pdf
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

BASE_DIR = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(BASE_DIR))


def parse_args():
    parser = argparse.ArgumentParser(prog="merge_texts_v2.py", description="Merge multiple OCR text outputs.")
    input_group = parser.add_mutually_exclusive_group(required=True)
    input_group.add_argument("-i", "--inputs", nargs="+", help="Input .txt files")
    input_group.add_argument("--dir", help="Directory containing all .txt files to merge")

    parser.add_argument("-o", "--output", default="./merged.txt", help="Output file path")
    parser.add_argument("--format", choices=["txt", "json", "docx"], default="txt", help="Output format")
    parser.add_argument("--separator", default="\n\n---\n\n", help="Separator between texts")
    return parser.parse_args()


def main():
    args = parse_args()

    # Collect input files
    if args.inputs:
        text_files = [Path(f) for f in args.inputs]
    else:
        dir_path = Path(args.dir)
        text_files = sorted(dir_path.glob("*.txt"))

    all_texts = []
    file_infos = []

    for tf in text_files:
        if tf.exists():
            content = tf.read_text(encoding="utf-8")
            all_texts.append(content)
            file_infos.append({"filename": tf.name, "lines": len(content.strip().splitlines())})
        else:
            print(f"[WARN] Not found: {tf}")

    if not all_texts:
        print("[ERROR] No valid input files found.")
        sys.exit(1)

    merged = args.separator.join(all_texts)

    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)

    if args.format == "json":
        import json
        output.write_text(json.dumps({
            "merged_text": merged,
            "source_files": file_infos,
            "total_lines": sum(f["lines"] for f in file_infos),
        }, indent=2, ensure_ascii=False), encoding="utf-8")
    elif args.format == "docx":
        from docx import Document
        doc = Document()
        doc.add_heading("Merged OCR Output", level=1)
        doc.add_paragraph(f"Source files: {len(file_infos)}", style="Intense Quote")
        for fi in file_infos:
            doc.add_heading(fi["filename"], level=2)
            doc.add_paragraph(fi["merged_text"])
        from io import BytesIO as BIO
        buf = BIO()
        doc.save(buf)
        output.write_bytes(buf.getvalue())
    else:
        output.write_text(merged, encoding="utf-8")

    print(f"Merged {len(file_infos)} files ({sum(f['lines'] for f in file_infos)} lines) → {output}")


if __name__ == "__main__":
    main()
